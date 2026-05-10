"""Unit tests for DOCXConverter — Task 5.3."""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Pt

from document2markdown.converter_docx import DOCXConverter
from document2markdown.document_model import (
    HeadingBlock,
    LinkBlock,
    ListBlock,
    ParagraphBlock,
    TableBlock,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx(tmp_path: Path, builder) -> Path:
    doc = DocxDocument()
    builder(doc)
    p = tmp_path / "test.docx"
    doc.save(str(p))
    return p


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestHeadings:
    def test_h1(self, tmp_path):
        p = _make_docx(tmp_path, lambda d: d.add_heading("Title One", level=1))
        result = DOCXConverter().convert(p)
        headings = [b for b in result.blocks if isinstance(b, HeadingBlock)]
        assert any(h.level == 1 and h.text == "Title One" for h in headings)

    def test_h2(self, tmp_path):
        p = _make_docx(tmp_path, lambda d: d.add_heading("Section", level=2))
        result = DOCXConverter().convert(p)
        headings = [b for b in result.blocks if isinstance(b, HeadingBlock)]
        assert any(h.level == 2 and h.text == "Section" for h in headings)

    def test_h3(self, tmp_path):
        p = _make_docx(tmp_path, lambda d: d.add_heading("Subsection", level=3))
        result = DOCXConverter().convert(p)
        headings = [b for b in result.blocks if isinstance(b, HeadingBlock)]
        assert any(h.level == 3 and h.text == "Subsection" for h in headings)


class TestLists:
    def test_unordered_list(self, tmp_path):
        def build(doc):
            doc.add_paragraph("Apple", style="List Bullet")
            doc.add_paragraph("Banana", style="List Bullet")

        p = _make_docx(tmp_path, build)
        result = DOCXConverter().convert(p)
        lists = [b for b in result.blocks if isinstance(b, ListBlock)]
        assert lists, "Expected at least one ListBlock"
        assert any(not lb.ordered for lb in lists)
        items = [i for lb in lists if not lb.ordered for i in lb.items]
        assert "Apple" in items
        assert "Banana" in items

    def test_ordered_list(self, tmp_path):
        def build(doc):
            doc.add_paragraph("First", style="List Number")
            doc.add_paragraph("Second", style="List Number")

        p = _make_docx(tmp_path, build)
        result = DOCXConverter().convert(p)
        lists = [b for b in result.blocks if isinstance(b, ListBlock)]
        assert lists, "Expected at least one ListBlock"
        assert any(lb.ordered for lb in lists)
        items = [i for lb in lists if lb.ordered for i in lb.items]
        assert "First" in items
        assert "Second" in items


class TestTable:
    def test_table_headers_and_rows(self, tmp_path):
        def build(doc):
            tbl = doc.add_table(rows=3, cols=2)
            tbl.cell(0, 0).text = "Name"
            tbl.cell(0, 1).text = "Age"
            tbl.cell(1, 0).text = "Alice"
            tbl.cell(1, 1).text = "30"
            tbl.cell(2, 0).text = "Bob"
            tbl.cell(2, 1).text = "25"

        p = _make_docx(tmp_path, build)
        result = DOCXConverter().convert(p)
        tables = [b for b in result.blocks if isinstance(b, TableBlock)]
        assert tables, "Expected at least one TableBlock"
        tb = tables[0]
        assert tb.headers == ["Name", "Age"]
        assert ["Alice", "30"] in tb.rows
        assert ["Bob", "25"] in tb.rows


class TestHyperlink:
    def test_hyperlink_produces_link_block(self, tmp_path):
        """A paragraph containing a hyperlink should produce a LinkBlock."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = DocxDocument()
        para = doc.add_paragraph()

        # Add the hyperlink relationship via the proper python-docx API so
        # the document can be saved without errors.
        r_id = para.part.relate_to(
            "https://example.com",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        # Build the hyperlink XML element referencing that relationship.
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "Click here"
        run.append(t)
        hyperlink.append(run)
        para._element.append(hyperlink)

        p = tmp_path / "link.docx"
        doc.save(str(p))

        result = DOCXConverter().convert(p)
        links = [b for b in result.blocks if isinstance(b, LinkBlock)]
        assert links, "Expected at least one LinkBlock"
        assert links[0].text == "Click here"
        assert links[0].url == "https://example.com"


class TestParagraph:
    def test_plain_paragraph(self, tmp_path):
        p = _make_docx(tmp_path, lambda d: d.add_paragraph("Hello world."))
        result = DOCXConverter().convert(p)
        paras = [b for b in result.blocks if isinstance(b, ParagraphBlock)]
        assert any(b.text == "Hello world." for b in paras)



# ---------------------------------------------------------------------------
# _ListAccumulator
# ---------------------------------------------------------------------------

class TestListAccumulator:
    def test_flush_returns_none_when_empty(self):
        from document2markdown.converter_docx import _ListAccumulator
        acc = _ListAccumulator()
        assert acc.flush() is None

    def test_active_false_when_empty(self):
        from document2markdown.converter_docx import _ListAccumulator
        acc = _ListAccumulator()
        assert acc.active is False

    def test_active_true_after_add(self):
        from document2markdown.converter_docx import _ListAccumulator
        acc = _ListAccumulator()
        acc.add("item", ordered=False)
        assert acc.active is True

    def test_flush_clears_items(self):
        from document2markdown.converter_docx import _ListAccumulator
        acc = _ListAccumulator()
        acc.add("a", ordered=True)
        acc.add("b", ordered=True)
        block = acc.flush()
        assert block is not None
        assert block.items == ["a", "b"]
        assert block.ordered is True
        assert acc.active is False


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

class TestStyleHelpers:
    def test_heading_level_returns_none_for_normal(self):
        from document2markdown.converter_docx import _heading_level
        assert _heading_level("Normal") is None

    def test_heading_level_returns_correct_level(self):
        from document2markdown.converter_docx import _heading_level
        for i in range(1, 7):
            assert _heading_level(f"Heading {i}") == i

    def test_is_code_style_true(self):
        from document2markdown.converter_docx import _is_code_style
        assert _is_code_style("Code Block") is True
        assert _is_code_style("code") is True

    def test_is_code_style_false(self):
        from document2markdown.converter_docx import _is_code_style
        assert _is_code_style("Normal") is False

    def test_content_type_to_ext_unknown(self):
        from document2markdown.converter_docx import _content_type_to_ext
        assert _content_type_to_ext("application/octet-stream") == ".bin"

    def test_content_type_to_vector_format_emf(self):
        from document2markdown.converter_docx import _content_type_to_vector_format
        assert _content_type_to_vector_format("image/x-emf") == "emf"

    def test_content_type_to_vector_format_none(self):
        from document2markdown.converter_docx import _content_type_to_vector_format
        assert _content_type_to_vector_format("image/png") is None


# ---------------------------------------------------------------------------
# Empty paragraph skipped
# ---------------------------------------------------------------------------

class TestEmptyParagraph:
    def test_empty_paragraph_produces_no_block(self, tmp_path):
        def build(doc):
            doc.add_paragraph("")
            doc.add_paragraph("Real content")

        p = _make_docx(tmp_path, build)
        result = DOCXConverter().convert(p)
        paras = [b for b in result.blocks if isinstance(b, ParagraphBlock)]
        texts = [b.text for b in paras]
        assert "Real content" in texts
        assert "" not in texts


# ---------------------------------------------------------------------------
# Code style paragraph
# ---------------------------------------------------------------------------

class TestCodeStyle:
    def test_code_style_paragraph(self, tmp_path):
        from document2markdown.document_model import CodeBlock

        def build(doc):
            # python-docx won't have a "Code Block" style by default,
            # so we use a monospace font run to trigger _run_has_code_style
            para = doc.add_paragraph()
            run = para.add_run("print('hello')")
            run.font.name = "Courier New"

        p = _make_docx(tmp_path, build)
        result = DOCXConverter().convert(p)
        code_blocks = [b for b in result.blocks if isinstance(b, CodeBlock)]
        assert len(code_blocks) == 1
        assert "print" in code_blocks[0].text


# ---------------------------------------------------------------------------
# Table edge cases
# ---------------------------------------------------------------------------

class TestTableEdgeCases:
    def test_empty_table_produces_empty_table_block(self, tmp_path):
        """A table with no rows should produce a TableBlock with empty headers/rows."""
        from document2markdown.converter_docx import DOCXConverter
        from document2markdown.document_model import TableBlock
        from unittest.mock import MagicMock

        converter = DOCXConverter()
        tbl = MagicMock()
        tbl.rows = []
        block = converter._convert_table(tbl)
        assert isinstance(block, TableBlock)
        assert block.headers == []
        assert block.rows == []

    def test_single_row_table_uses_row_as_headers(self, tmp_path):
        """A table with only one row should use it as headers with no data rows."""
        def build(doc):
            tbl = doc.add_table(rows=1, cols=2)
            tbl.cell(0, 0).text = "Col A"
            tbl.cell(0, 1).text = "Col B"

        p = _make_docx(tmp_path, build)
        result = DOCXConverter().convert(p)
        tables = [b for b in result.blocks if isinstance(b, TableBlock)]
        assert tables
        assert tables[0].headers == ["Col A", "Col B"]
        assert tables[0].rows == []


# ---------------------------------------------------------------------------
# Trailing list flush
# ---------------------------------------------------------------------------

class TestTrailingListFlush:
    def test_list_at_end_of_document_is_flushed(self, tmp_path):
        """A list at the very end of the document (no following paragraph) should be emitted."""
        def build(doc):
            doc.add_paragraph("Intro paragraph.")
            doc.add_paragraph("Item one", style="List Bullet")
            doc.add_paragraph("Item two", style="List Bullet")
            # No paragraph after — tests trailing flush

        p = _make_docx(tmp_path, build)
        result = DOCXConverter().convert(p)
        lists = [b for b in result.blocks if isinstance(b, ListBlock)]
        assert lists
        assert "Item one" in lists[0].items
        assert "Item two" in lists[0].items


# ---------------------------------------------------------------------------
# Mixed list types flush on type change
# ---------------------------------------------------------------------------

class TestMixedListTypes:
    def test_bullet_then_numbered_produces_two_list_blocks(self, tmp_path):
        """Switching from bullet to numbered list should flush and start a new ListBlock."""
        def build(doc):
            doc.add_paragraph("Bullet A", style="List Bullet")
            doc.add_paragraph("Bullet B", style="List Bullet")
            doc.add_paragraph("Number 1", style="List Number")
            doc.add_paragraph("Number 2", style="List Number")

        p = _make_docx(tmp_path, build)
        result = DOCXConverter().convert(p)
        lists = [b for b in result.blocks if isinstance(b, ListBlock)]
        assert len(lists) == 2
        assert lists[0].ordered is False
        assert lists[1].ordered is True


# ---------------------------------------------------------------------------
# _try_extract_hyperlink — no hyperlink in paragraph
# ---------------------------------------------------------------------------

class TestHyperlinkEdgeCases:
    def test_paragraph_without_hyperlink_returns_none(self, tmp_path):
        """_try_extract_hyperlink should return None for a plain paragraph."""
        p = _make_docx(tmp_path, lambda d: d.add_paragraph("No link here."))
        from docx import Document as DocxDocument
        from docx.text.paragraph import Paragraph
        doc = DocxDocument(str(p))
        para = doc.paragraphs[0]
        result = DOCXConverter()._try_extract_hyperlink(para, doc)
        assert result is None


# ---------------------------------------------------------------------------
# _extract_raster_image — missing relationship
# ---------------------------------------------------------------------------

class TestExtractRasterImageMissing:
    def test_missing_rel_adds_warning_and_returns_none(self, tmp_path):
        """If the image relationship is not found, a warning is added and None returned."""
        from unittest.mock import MagicMock
        from docx.text.paragraph import Paragraph

        converter = DOCXConverter()
        embedded = []
        warnings = []

        para = MagicMock(spec=Paragraph)
        para.part.rels.get.return_value = None

        result = converter._extract_raster_image("rId99", para, embedded, warnings)
        assert result is None
        assert any("rId99" in w for w in warnings)


# ---------------------------------------------------------------------------
# _convert_vector_asset — VectorConversionError
# ---------------------------------------------------------------------------

class TestConvertVectorAssetFailure:
    def test_vector_conversion_error_adds_warning(self, tmp_path):
        """VectorConversionError during vector conversion should add a warning and return None."""
        from unittest.mock import patch
        from document2markdown.converter_vector import VectorConversionError

        converter = DOCXConverter()
        embedded = []
        warnings = []

        with patch.object(converter._vector, "convert", side_effect=VectorConversionError("no tool")):
            result = converter._convert_vector_asset(b"fake", "emf", embedded, warnings)

        assert result is None
        assert any("emf" in w for w in warnings)


# ---------------------------------------------------------------------------
# _style_name exception path
# ---------------------------------------------------------------------------

class TestStyleNameException:
    def test_style_name_returns_empty_on_exception(self):
        """_style_name should return '' if para.style raises."""
        from unittest.mock import MagicMock, PropertyMock
        from document2markdown.converter_docx import _style_name
        para = MagicMock()
        type(para).style = PropertyMock(side_effect=Exception("no style"))
        assert _style_name(para) == ""


# ---------------------------------------------------------------------------
# _is_ordered_list — numbering_part is None
# ---------------------------------------------------------------------------

class TestIsOrderedListNullNumberingPart:
    def test_returns_false_when_numbering_part_is_none(self, tmp_path):
        """_is_ordered_list should return False when doc has no numbering part."""
        from unittest.mock import MagicMock
        from docx.oxml.ns import qn
        from lxml import etree

        converter = DOCXConverter()
        doc = MagicMock()
        doc.part.numbering_part = None

        # Build a minimal num_id_el
        num_id_el = etree.Element(qn("w:numId"))
        num_id_el.set(qn("w:val"), "1")

        result = converter._is_ordered_list(doc, num_id_el)
        assert result is False

    def test_returns_false_when_num_id_el_is_none(self):
        from document2markdown.converter_docx import DOCXConverter
        from unittest.mock import MagicMock
        converter = DOCXConverter()
        result = converter._is_ordered_list(MagicMock(), None)
        assert result is False


# ---------------------------------------------------------------------------
# _extract_vector_by_rel — non-vector content type (raster fallback)
# ---------------------------------------------------------------------------

class TestExtractVectorByRelRasterFallback:
    def test_non_vector_content_type_treated_as_raster(self):
        """If the rel target is not a vector format, it should be stored as raster."""
        from unittest.mock import MagicMock
        from document2markdown.document_model import ImageBlock

        converter = DOCXConverter()
        embedded = []
        warnings = []

        target_part = MagicMock()
        target_part.blob = b"\x89PNG\r\n"
        target_part.content_type = "image/png"
        target_part.partname = "/word/media/image1.png"

        rel = MagicMock()
        rel.target_part = target_part

        para = MagicMock()
        para.part.rels.get.return_value = rel

        result = converter._extract_vector_by_rel("rId1", para, embedded, warnings)
        assert isinstance(result, ImageBlock)
        assert embedded[0].extension == ".png"
        assert warnings == []

    def test_missing_vector_rel_adds_warning(self):
        """If the vector relationship is not found, a warning is added."""
        from unittest.mock import MagicMock

        converter = DOCXConverter()
        warnings = []
        para = MagicMock()
        para.part.rels.get.return_value = None

        result = converter._extract_vector_by_rel("rId99", para, [], warnings)
        assert result is None
        assert any("rId99" in w for w in warnings)

    def test_exception_in_vector_rel_adds_warning(self):
        """If accessing the rel raises, a warning is added."""
        from unittest.mock import MagicMock

        converter = DOCXConverter()
        warnings = []
        para = MagicMock()
        para.part.rels.get.side_effect = Exception("rel error")

        result = converter._extract_vector_by_rel("rId5", para, [], warnings)
        assert result is None
        assert any("rId5" in w for w in warnings)


# ---------------------------------------------------------------------------
# _extract_raster_image — exception path
# ---------------------------------------------------------------------------

class TestExtractRasterImageException:
    def test_exception_adds_warning_and_returns_none(self):
        """If accessing image_part raises, a warning is added."""
        from unittest.mock import MagicMock

        converter = DOCXConverter()
        warnings = []
        para = MagicMock()
        rel = MagicMock()
        rel.target_part = MagicMock()
        type(rel.target_part).blob = property(lambda self: (_ for _ in ()).throw(Exception("blob error")))
        para.part.rels.get.return_value = rel

        result = converter._extract_raster_image("rId1", para, [], warnings)
        assert result is None
        assert any("rId1" in w for w in warnings)
