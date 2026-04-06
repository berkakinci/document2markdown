"""DOCX converter — maps python-docx document elements to IR blocks.

Supported mappings:

* Heading 1–6 paragraph styles → :class:`HeadingBlock`
* Normal / body paragraphs     → :class:`ParagraphBlock`
* Ordered / unordered lists    → :class:`ListBlock`
* Tables                       → :class:`TableBlock`
* Hyperlinks                   → :class:`LinkBlock`
* Code-styled runs             → :class:`CodeBlock`
* Raster images                → :class:`EmbeddedAsset` + :class:`ImageBlock`
* EMF / WMF drawing objects    → :class:`VectorConverter` → :class:`EmbeddedAsset` + :class:`ImageBlock`
* Unrenderable elements        → :class:`UnsupportedBlock`
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import TYPE_CHECKING
from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from document2markdown.config import RASTER_DPI
from document2markdown.converter_base import BaseConverter
from document2markdown.converter_vector import VectorConverter, VectorConversionError
from document2markdown.document_model import (
    CodeBlock,
    ConversionResult,
    EmbeddedAsset,
    HeadingBlock,
    ImageBlock,
    LinkBlock,
    ListBlock,
    ParagraphBlock,
    TableBlock,
    UnsupportedBlock,
)

if TYPE_CHECKING:
    from document2markdown.document_model import Block

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Style-name → heading level mapping
# ---------------------------------------------------------------------------

# python-docx exposes style names like "Heading 1", "Heading 2", …
_HEADING_RE = re.compile(r"^[Hh]eading\s+([1-6])$")

# Style names that indicate code / monospace content
_CODE_STYLE_NAMES = {
    "code",
    "code block",
    "codeblock",
    "pre",
    "preformatted",
    "verbatim",
    "monospace",
    "html code",
    "html preformatted",
}

# List style name patterns
_BULLET_LIST_RE = re.compile(r"list\s*bullet", re.IGNORECASE)
_NUMBER_LIST_RE = re.compile(r"list\s*number", re.IGNORECASE)

# Relationship type for hyperlinks
_HYPERLINK_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)

# XML namespaces used in DOCX
_NS = {
    "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r":   "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "v":   "urn:schemas-microsoft-com:vml",
    "o":   "urn:schemas-microsoft-com:office:office",
}


def _style_name(para: Paragraph) -> str:
    """Return the normalised style name for *para*, or empty string."""
    try:
        return (para.style.name or "").strip()
    except Exception:  # noqa: BLE001
        return ""


def _heading_level(style: str) -> int | None:
    """Return 1–6 if *style* is a heading style, else None."""
    m = _HEADING_RE.match(style)
    return int(m.group(1)) if m else None


def _is_code_style(style: str) -> bool:
    return style.lower() in _CODE_STYLE_NAMES


def _para_text(para: Paragraph) -> str:
    """Return the full text of *para* (all runs concatenated)."""
    return para.text


def _run_has_code_style(para: Paragraph) -> bool:
    """Return True if every non-empty run in *para* uses a monospace font."""
    for run in para.runs:
        if not run.text.strip():
            continue
        font_name = (run.font.name or "").lower()
        if "courier" in font_name or "mono" in font_name or "consolas" in font_name:
            return True
    return False


# ---------------------------------------------------------------------------
# List accumulator helpers
# ---------------------------------------------------------------------------

class _ListAccumulator:
    """Accumulates consecutive list paragraphs into a single ListBlock."""

    def __init__(self) -> None:
        self.ordered: bool = False
        self.items: list[str] = []

    def flush(self) -> ListBlock | None:
        if not self.items:
            return None
        block = ListBlock(ordered=self.ordered, items=list(self.items))
        self.items.clear()
        return block

    def add(self, text: str, ordered: bool) -> None:
        if self.items and self.ordered != ordered:
            # Type changed — caller should flush first
            pass
        self.ordered = ordered
        self.items.append(text)

    @property
    def active(self) -> bool:
        return bool(self.items)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

class DOCXConverter(BaseConverter):
    """Convert a ``.docx`` file to a :class:`ConversionResult`.

    Parameters
    ----------
    raster_dpi:
        DPI used when rasterizing vector graphics as a last resort.
    """

    def __init__(self, raster_dpi: int = RASTER_DPI) -> None:
        self._vector = VectorConverter(raster_dpi=raster_dpi)

    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------

    def convert(self, source_path: Path) -> ConversionResult:
        """Convert *source_path* (.docx) to a :class:`ConversionResult`."""
        doc = DocxDocument(str(source_path))

        blocks: list[Block] = []
        embedded: list[EmbeddedAsset] = []
        warnings: list[str] = []

        # Open the underlying ZIP for raw image/drawing extraction
        with ZipFile(str(source_path), "r") as zf:
            zip_names = set(zf.namelist())

            list_acc = _ListAccumulator()

            # Iterate top-level body elements in document order.
            # python-docx exposes doc.paragraphs and doc.tables but not in
            # interleaved order.  We walk doc.element.body directly.
            for child in doc.element.body:
                tag = child.tag

                # ---- Paragraph ----
                if tag == qn("w:p"):
                    para = Paragraph(child, doc)
                    para_blocks, flush_needed = self._convert_paragraph(
                        para, doc, zf, zip_names, embedded, warnings
                    )
                    if flush_needed and list_acc.active:
                        lb = list_acc.flush()
                        if lb:
                            blocks.append(lb)
                    for b in para_blocks:
                        if isinstance(b, _ListItem):
                            list_acc.add(b.text, b.ordered)
                        else:
                            if list_acc.active:
                                lb = list_acc.flush()
                                if lb:
                                    blocks.append(lb)
                            blocks.append(b)
                    continue

                # ---- Table ----
                if tag == qn("w:tbl"):
                    if list_acc.active:
                        lb = list_acc.flush()
                        if lb:
                            blocks.append(lb)
                    tbl = Table(child, doc)
                    blocks.append(self._convert_table(tbl))
                    continue

                # ---- Section properties (ignored) ----
                if tag == qn("w:sectPr"):
                    continue

                # ---- Anything else ----
                desc = f"unsupported body element <{child.tag}>"
                warnings.append(desc)
                blocks.append(UnsupportedBlock(description=desc))

            # Flush any trailing list
            if list_acc.active:
                lb = list_acc.flush()
                if lb:
                    blocks.append(lb)

        return ConversionResult(blocks=blocks, embedded=embedded, warnings=warnings)

    # ------------------------------------------------------------------
    # Paragraph conversion
    # ------------------------------------------------------------------

    def _convert_paragraph(
        self,
        para: Paragraph,
        doc: DocxDocument,
        zf: ZipFile,
        zip_names: set[str],
        embedded: list[EmbeddedAsset],
        warnings: list[str],
    ) -> tuple[list[Block], bool]:
        """Convert a single paragraph.

        Returns ``(blocks, flush_needed)`` where *flush_needed* signals that
        any pending list accumulator should be flushed before appending.
        """
        style = _style_name(para)
        text = _para_text(para).strip()

        # --- Heading ---
        level = _heading_level(style)
        if level is not None:
            return [HeadingBlock(level=level, text=text)], True

        # --- Code style ---
        if _is_code_style(style) or _run_has_code_style(para):
            return [CodeBlock(language=None, text=text)], True

        # --- List (bullet) ---
        if _BULLET_LIST_RE.search(style):
            return [_ListItem(text=text, ordered=False)], False

        # --- List (numbered) ---
        if _NUMBER_LIST_RE.search(style):
            return [_ListItem(text=text, ordered=True)], False

        # --- Check for list via numPr XML ---
        num_pr = para._element.find(qn("w:numPr"))  # noqa: SLF001
        if num_pr is not None:
            ilvl_el = num_pr.find(qn("w:ilvl"))
            num_id_el = num_pr.find(qn("w:numId"))
            ordered = self._is_ordered_list(doc, num_id_el)
            return [_ListItem(text=text, ordered=ordered)], False

        # --- Hyperlink (paragraph contains only a hyperlink run) ---
        link_block = self._try_extract_hyperlink(para, doc)
        if link_block is not None:
            return [link_block], True

        # --- Inline images / drawings ---
        image_blocks = self._extract_images(para, doc, zf, zip_names, embedded, warnings)
        if image_blocks:
            result: list[Block] = []
            if text:
                result.append(ParagraphBlock(text=text))
            result.extend(image_blocks)
            return result, True

        # --- Plain paragraph ---
        if text:
            return [ParagraphBlock(text=text)], True

        # Empty paragraph — skip
        return [], False

    # ------------------------------------------------------------------
    # List helpers
    # ------------------------------------------------------------------

    def _is_ordered_list(self, doc: DocxDocument, num_id_el) -> bool:
        """Determine if a list numId refers to an ordered (numbered) list."""
        if num_id_el is None:
            return False
        try:
            num_id = int(num_id_el.get(qn("w:val"), "0"))
            # Walk numbering.xml to find the abstract num format
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                return False
            numbering_el = numbering_part._element  # noqa: SLF001
            # Find <w:num w:numId="...">
            for num_el in numbering_el.findall(qn("w:num")):
                if int(num_el.get(qn("w:numId"), "-1")) == num_id:
                    abs_ref = num_el.find(qn("w:abstractNumId"))
                    if abs_ref is None:
                        return False
                    abs_id = int(abs_ref.get(qn("w:val"), "-1"))
                    # Find <w:abstractNum w:abstractNumId="...">
                    for abs_el in numbering_el.findall(qn("w:abstractNum")):
                        if int(abs_el.get(qn("w:abstractNumId"), "-1")) == abs_id:
                            # Check first level's numFmt
                            lvl = abs_el.find(qn("w:lvl"))
                            if lvl is not None:
                                fmt_el = lvl.find(qn("w:numFmt"))
                                if fmt_el is not None:
                                    fmt = fmt_el.get(qn("w:val"), "")
                                    return fmt not in ("bullet", "none", "")
        except Exception as exc:  # noqa: BLE001
            logger.debug("Could not determine list type: %s", exc)
        return False

    # ------------------------------------------------------------------
    # Hyperlink extraction
    # ------------------------------------------------------------------

    def _try_extract_hyperlink(
        self, para: Paragraph, doc: DocxDocument
    ) -> LinkBlock | None:
        """If *para* contains a hyperlink, return a :class:`LinkBlock`."""
        hyperlinks = para._element.findall(qn("w:hyperlink"))  # noqa: SLF001
        if not hyperlinks:
            return None

        # Use the first hyperlink found
        hl = hyperlinks[0]
        r_id = hl.get(qn("r:id"))
        url = ""
        if r_id:
            try:
                rel = para.part.rels.get(r_id)
                if rel and rel.reltype == _HYPERLINK_REL_TYPE:
                    url = rel.target_ref
            except Exception:  # noqa: BLE001
                pass

        # Collect text from all runs inside the hyperlink
        texts = []
        for run_el in hl.findall(qn("w:r")):
            t_el = run_el.find(qn("w:t"))
            if t_el is not None and t_el.text:
                texts.append(t_el.text)
        link_text = "".join(texts).strip() or para.text.strip()

        if url or link_text:
            return LinkBlock(text=link_text, url=url)
        return None

    # ------------------------------------------------------------------
    # Image / drawing extraction
    # ------------------------------------------------------------------

    def _extract_images(
        self,
        para: Paragraph,
        doc: DocxDocument,
        zf: ZipFile,
        zip_names: set[str],
        embedded: list[EmbeddedAsset],
        warnings: list[str],
    ) -> list[ImageBlock]:
        """Extract all images/drawings from *para* and return ImageBlocks."""
        blocks: list[ImageBlock] = []
        para_el = para._element  # noqa: SLF001

        # --- DrawingML images (modern DOCX) ---
        for drawing in para_el.iter(qn("w:drawing")):
            blip_fills = drawing.iter(qn("a:blip"))
            for blip in blip_fills:
                r_id = blip.get(qn("r:embed"))
                if not r_id:
                    continue
                block = self._extract_raster_image(r_id, para, embedded, warnings)
                if block:
                    blocks.append(block)

        # --- VML images (older DOCX / compatibility) ---
        for pict in para_el.iter(qn("w:pict")):
            # VML imagedata
            for imgdata in pict.iter("{urn:schemas-microsoft-com:vml}imagedata"):
                r_id = imgdata.get(qn("r:id"))
                if not r_id:
                    continue
                block = self._extract_raster_image(r_id, para, embedded, warnings)
                if block:
                    blocks.append(block)

            # VML shape with EMF/WMF (OLE drawing)
            for shape in pict.iter("{urn:schemas-microsoft-com:vml}shape"):
                block = self._extract_vml_vector(shape, pict, para, zf, zip_names, embedded, warnings)
                if block:
                    blocks.append(block)

        # --- mc:AlternateContent (may contain EMF/WMF fallback) ---
        for alt_content in para_el.iter(
            "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"
        ):
            block = self._extract_alternate_content(alt_content, para, zf, zip_names, embedded, warnings)
            if block:
                blocks.append(block)

        return blocks

    def _extract_raster_image(
        self,
        r_id: str,
        para: Paragraph,
        embedded: list[EmbeddedAsset],
        warnings: list[str],
    ) -> ImageBlock | None:
        """Extract a raster image by relationship ID."""
        try:
            rel = para.part.rels.get(r_id)
            if rel is None:
                warnings.append(f"image relationship {r_id!r} not found")
                return None
            image_part = rel.target_part
            data: bytes = image_part.blob
            content_type: str = image_part.content_type or ""
            ext = _content_type_to_ext(content_type)
            original_name = Path(image_part.partname).name
            alt = original_name or "image"
            asset = EmbeddedAsset(
                data=data,
                extension=ext,
                original_name=original_name,
                alt_text=alt,
                source_vector_format=None,
            )
            idx = len(embedded)
            embedded.append(asset)
            return ImageBlock(asset_index=idx, alt=alt)
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"failed to extract raster image {r_id!r}: {exc}")
            return None

    def _extract_vml_vector(
        self,
        shape,
        pict,
        para: Paragraph,
        zf: ZipFile,
        zip_names: set[str],
        embedded: list[EmbeddedAsset],
        warnings: list[str],
    ) -> ImageBlock | None:
        """Try to extract an EMF/WMF from a VML <v:shape> element."""
        # OLE objects store the vector data in a separate part referenced by
        # <o:OLEObject r:id="..."> or via imagedata
        ole_ns = "urn:schemas-microsoft-com:office:office"
        for ole in pict.iter(f"{{{ole_ns}}}OLEObject"):
            r_id = ole.get(qn("r:id"))
            if r_id:
                return self._extract_vector_by_rel(r_id, para, embedded, warnings)
        return None

    def _extract_alternate_content(
        self,
        alt_content,
        para: Paragraph,
        zf: ZipFile,
        zip_names: set[str],
        embedded: list[EmbeddedAsset],
        warnings: list[str],
    ) -> ImageBlock | None:
        """Extract image from mc:AlternateContent, preferring the Fallback."""
        mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
        # Try Fallback first (often contains EMF/WMF for older readers)
        fallback = alt_content.find(f"{{{mc_ns}}}Fallback")
        choice = alt_content.find(f"{{{mc_ns}}}Choice")

        for container in [fallback, choice]:
            if container is None:
                continue
            # DrawingML blip inside
            for blip in container.iter(qn("a:blip")):
                r_id = blip.get(qn("r:embed"))
                if r_id:
                    return self._extract_raster_image(r_id, para, embedded, warnings)
        return None

    def _extract_vector_by_rel(
        self,
        r_id: str,
        para: Paragraph,
        embedded: list[EmbeddedAsset],
        warnings: list[str],
    ) -> ImageBlock | None:
        """Extract a vector graphic (EMF/WMF) by relationship ID."""
        try:
            rel = para.part.rels.get(r_id)
            if rel is None:
                warnings.append(f"vector relationship {r_id!r} not found")
                return None
            target_part = rel.target_part
            data: bytes = target_part.blob
            content_type: str = target_part.content_type or ""
            src_fmt = _content_type_to_vector_format(content_type)
            if src_fmt is None:
                # Not a known vector format — treat as raster
                ext = _content_type_to_ext(content_type)
                original_name = Path(target_part.partname).name
                asset = EmbeddedAsset(
                    data=data,
                    extension=ext,
                    original_name=original_name,
                    alt_text=original_name or "image",
                    source_vector_format=None,
                )
                idx = len(embedded)
                embedded.append(asset)
                return ImageBlock(asset_index=idx, alt=asset.alt_text)

            # Pass through VectorConverter
            return self._convert_vector_asset(data, src_fmt, embedded, warnings)
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"failed to extract vector {r_id!r}: {exc}")
            return None

    def _convert_vector_asset(
        self,
        data: bytes,
        src_fmt: str,
        embedded: list[EmbeddedAsset],
        warnings: list[str],
    ) -> ImageBlock | None:
        """Run *data* through :class:`VectorConverter` and emit an asset."""
        try:
            out_bytes, ext = self._vector.convert(data, src_fmt)  # type: ignore[arg-type]
            asset = EmbeddedAsset(
                data=out_bytes,
                extension=ext,
                original_name=None,
                alt_text=f"{src_fmt} drawing",
                source_vector_format=src_fmt,
            )
            idx = len(embedded)
            embedded.append(asset)
            return ImageBlock(asset_index=idx, alt=asset.alt_text)
        except VectorConversionError as exc:
            warnings.append(f"vector conversion failed ({src_fmt}): {exc}")
            return None

    # ------------------------------------------------------------------
    # Table conversion
    # ------------------------------------------------------------------

    def _convert_table(self, tbl: Table) -> TableBlock:
        """Convert a python-docx :class:`Table` to a :class:`TableBlock`."""
        rows_data: list[list[str]] = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return TableBlock(headers=[], rows=[])

        headers = rows_data[0]
        rows = rows_data[1:]
        return TableBlock(headers=headers, rows=rows)


# ---------------------------------------------------------------------------
# Internal sentinel for list items (not a public Block type)
# ---------------------------------------------------------------------------

class _ListItem:
    """Temporary sentinel used during paragraph accumulation."""

    def __init__(self, text: str, ordered: bool) -> None:
        self.text = text
        self.ordered = ordered


# ---------------------------------------------------------------------------
# MIME / extension helpers
# ---------------------------------------------------------------------------

_CONTENT_TYPE_TO_EXT: dict[str, str] = {
    "image/png":  ".png",
    "image/jpeg": ".jpg",
    "image/jpg":  ".jpg",
    "image/gif":  ".gif",
    "image/bmp":  ".bmp",
    "image/tiff": ".tif",
    "image/webp": ".webp",
    "image/svg+xml": ".svg",
    "image/x-emf": ".emf",
    "image/x-wmf": ".wmf",
    "image/emf":   ".emf",
    "image/wmf":   ".wmf",
}

_VECTOR_CONTENT_TYPES: dict[str, str] = {
    "image/x-emf": "emf",
    "image/emf":   "emf",
    "image/x-wmf": "wmf",
    "image/wmf":   "wmf",
}


def _content_type_to_ext(content_type: str) -> str:
    """Map a MIME content type to a file extension."""
    ct = content_type.split(";")[0].strip().lower()
    return _CONTENT_TYPE_TO_EXT.get(ct, ".bin")


def _content_type_to_vector_format(content_type: str) -> str | None:
    """Return ``"emf"`` or ``"wmf"`` if *content_type* is a vector format, else None."""
    ct = content_type.split(";")[0].strip().lower()
    return _VECTOR_CONTENT_TYPES.get(ct)
